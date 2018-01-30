# -*-coding: utf-8-*-

import falcon, json, sys
from wsgiref import simple_server
import pymysql
from login_api import *
from config import *
import gunicorn
from query_project import *
from docx import Document
from docx.shared import Inches
from query_sample_info import *
sys.getdefaultencoding()
import datetime
from addinfo import templateword
from project_conf import *
from falcon_multipart.middleware import MultipartMiddleware
from flask import Flask, request, render_template as rt
from server import *

app = falcon.API(middleware=[MultipartMiddleware()])
# app = falcon.API()
app.req_options.auto_parse_form_urlencoded = True
app.req_options.keep_blank_qs_values = True


class Login():
    def on_post(self, req, resp):
        data = req.params
        print(data)
        #data = eval(list(data1.keys())[0])
        #print(data)
        """  firstName: '', lastName: '', passWord: '', email: '', liteLab: '', passWord2: '', email2: '', institutionName: """
        username = data["userName"]
        password = data["passWord"]
        db = Mysql(table_name="MYSQL")
        sql = """select * from register where email = \"%s\";""" % username
        data2 = db.fetch_one(sql)
        print(data2)
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        if data2:
            password_sql = data2[3]
            if password == password_sql:
                return_info = '{"success":true, "info": "Log in success!" }'
            else:
                return_info = '{"success":false, "info": "Wrong passWord!" }'
            resp.body = return_info
        else:
            resp.body = '{"success":false, "info": "Wrong userName!"}'
        db.commit()

class Register():
    def on_post(self, req, resp):
        data = req.params
        #data = eval(list(data1.keys())[0])
        """  firstName: '', lastName: '', passWord: '', email: '', liteLab: '', passWord2: '', email2: '', institutionName: """
        firstName = data["firstName"]
        lastName = data["lastName"]
        passWord = data["passWord"]
        email = data["email"]
        liteLab = data["liteLab"]
        passWord2 = data["passWord2"]
        email2 = data["email2"]
        institutionName = data["institutionName"]
        print(firstName)
        db = Mysql(table_name="MYSQL")
        sql = """select * from register where email = \"%s\";""" % email
        data2 = db.fetch_one(sql)
        print(data2)
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        if data2:
            return_info = {"success":False, "info":"%s exists!"%email }
            resp.body = json.dumps(return_info, ensure_ascii=False)
        else:
            sql = ' insert into register values(UUID(), \"%s\", \"%s\", \"%s\",\"%s\", \"%s\", \"%s\"); ' % (firstName, lastName, passWord, email, liteLab, institutionName)
            db.execute(sql)
            sql = "select * from register where passWord=\"%s\" and email=\"%s\";" % (passWord, email)
            out = db.fetch_one(sql)
            register_id = out[0]
            try:
                sql_login = "insert into login values(\"%s\",\"%s\", UUID(), \"%s\");" % (email, passWord, register_id)
                db.execute(sql_login)
            except Exception:
                sql_login_create = "create table if not exists login (ID VARCHAR(60), REGISTER_ID VARCHAR(60), userName VARCHAR(20), passWord VARCHAR(20))DEFAULT CHARSET = UTF8;"
                db.execute(sql_login_create)
                sql_login = "insert into login values(\"%s\",\"%s\", UUID(), \"%s\");" % (email, passWord, register_id)
                db.execute(sql_login)
            return_info = {"success":True, "info":data}
            resp.body = json.dumps(return_info, ensure_ascii=False)
        db.commit()

class query_sample(object):
    def on_post(self, req, resp):
        data = req.params
        project_info = str(data["project"])
        sampletype = str(data["sampleType"])
        samplename = str(data["sampleName"])
        username = data["userName"]
        # password = data["passWord"]
        size = int(data["size"])
        page = int(data['page'])
        print('size',size,'page',page)
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"

        sample_mysql = Mysql(table_name="PC_SAMPLE")
        register_mysql = Mysql(table_name="MYSQL")
        sql = "select * from register where email=\"%s\";" %  username
        out = register_mysql.fetch_one(sql)
        register_id = out[0]
        sql = "select * from sample where REGISTER_ID=\"%s\";" % register_id
        out_data = sample_mysql.fetch_all(sql)
        data = []
        if out_data:
            id = 0
            for out1 in out_data:
                sample_name = out1[2]
                sample_id = sample_name
                sample_detail = out1[3]
                project = out1[-5]
                sample_type = out1[-4]
                role = out1[-3]
                imported_by = out1[-2]
                imported_on = out1[-1]

                if project_info and samplename and sampletype:
                    if samplename == sample_id and project_info == project and sampletype == sample_type:
                        return_json = {"sampleName": sample_id, "sampleDetail": sample_detail, "project": project,
                                       "sampleType": sample_type, "role": role, "importedBy": imported_by,
                                       "importedOn": imported_on}
                    else:
                        continue
                elif project_info and samplename and not sampletype:
                    if project_info == project and sample_id==samplename:
                        return_json = {"sampleName": sample_id, "sampleDetail": sample_detail, "project": project,
                                       "sampleType": sample_type, "role": role, "importedBy": imported_by,
                                       "importedOn": imported_on}
                    else:
                        continue
                elif project_info and not samplename and not sampletype:
                    if project_info == project:
                        return_json = {"sampleName": sample_id, "sampleDetail": sample_detail, "project": project,
                                       "sampleType": sample_type, "role": role, "importedBy": imported_by,
                                       "importedOn": imported_on}
                    else:
                        continue
                elif not project_info and samplename and sampletype:
                    if sample_id == samplename and sampletype == sample_type:
                        return_json = {"sampleName": sample_id, "sampleDetail": sample_detail, "project": project,
                                       "sampleType": sample_type, "role": role, "importedBy": imported_by,
                                       "importedOn": imported_on}
                    else:
                        continue
                elif not project_info and not samplename and sampletype:
                    if sampletype == sample_type:
                        return_json = {"sampleName": sample_id, "sampleDetail": sample_detail, "project": project,
                                       "sampleType": sample_type, "role": role, "importedBy": imported_by,
                                       "importedOn": imported_on}
                    else:
                        continue
                elif not project_info and samplename and not sampletype:
                    if sample_id == samplename:
                        return_json = {"sampleName": sample_id, "sampleDetail": sample_detail, "project": project,
                                       "sampleType": sample_type, "role": role, "importedBy": imported_by,
                                        "importedOn": imported_on}
                    else:
                        continue
                elif project_info and not samplename and sampletype:
                    if project_info == project and sampletype == sample_type:
                        return_json = {"sampleName": sample_id, "sampleDetail": sample_detail, "project": project,
                                       "sampleType": sample_type, "role": role, "importedBy": imported_by,
                                       "importedOn": imported_on}
                    else:
                        continue
                elif not project_info and not samplename and not sampletype:
                    return_json = {"sampleName": sample_id, "sampleDetail": sample_detail, "project": project,
                                   "sampleType": sample_type, "role": role, "importedBy": imported_by,
                                   "importedOn": imported_on}

                data.append(return_json)
            range_tmp = range(len(data))
            mod = len(data) % size  # 取余
            reminder = int(len(data) / size)
            if mod != 0:
                if page == (reminder+1):
                    range_id = range_tmp[(size * (page - 1)):len(data)]
                else:
                    range_id = range_tmp[(size * (page - 1)):(size * (page))]
            else:
                range_id = range_tmp[(size * (page - 1)):(size * page)]
            out_json = []
            for index in range_id:
                out_json.append(data[index])
            length = str(len(out_json))
            total = str(len(data))
            return_info = '{"success":true, "info": %s, "length":%s, "total":%s}' % (json.dumps(out_json), length, total)
        else:
            return_info = '{"success":false, "info": "wrong!" }'
        sample_mysql.commit()
        register_mysql.commit()
        resp.body = return_info


class Download():
    def on_post(self, req, resp):
        params = req.params
        samplename = params["sampleName"]
        title = params["title"]
        if title:
            pass
        # if samplename:
        #     """先假设接收的参数 进行进一步的处理"""
        #     pass
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        document = Document()
        document.add_heading(u'普通的段落文字，这是 ', 0)

        p = document.add_paragraph(u'普通的段落文字，这是 ')
        p.add_run(u'加粗').bold = True
        p.add_run(u' 这是 ')
        p.add_run(u'斜体').italic = True

        document.add_heading(u'这是1级标题', level=1)
        document.add_paragraph(u'引用', style='IntenseQuote')

        document.add_paragraph(
            u'无序列表', style='ListBullet'
        )
        document.add_paragraph(
            u'有序列表', style='ListNumber'
        )

        # document.add_picture('qr.png', width=Inches(1.25))

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'

        recordset = [{'qty': 100, 'id': 15, 'desc': 'saksfkjsakj'}, {'qty': 99, 'id': 11, 'desc': '9fgduieef'},
                     {'qty': 87, 'id': 13, 'desc': 'uiiogdsagw'}, {'qty': 69, 'id': 14, 'desc': 'hgjhshsd'}]

        for item in recordset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['qty'])
            row_cells[1].text = str(item['id'])
            row_cells[2].text = item['desc']

        document.add_page_break()
        report_name = "%s.docx" % samplename
        # writer = csv.writer(response)
        document.save('/home/khl/web/dist/downloads/tmp/%s' % report_name)  # 生成文件放到该文件夹下面
        files = '/home/khl/web/dist/downloads/tmp/%s' % report_name
        if os.path.exists(files):
            result = {'success': True, 'info': 'http://192.168.1.144:8096/downloads/tmp/%s' % report_name, "reportName": report_name}
        else:
            result = {'success': False, 'info': '报告生成错误!'}
        print(result)
        resp.body = json.dumps(result, ensure_ascii=False)

class project_name(object):
    def on_post(self, req, resp):
        params = req.params
        type = params["type"]   # ngs 和 noNgs类型判断
        version = params["version"]   # 版本号判断
        db = Mysql(table_name="REPORT_MYSQL")
        sql = """select * from projectName;"""
        project_tmp = db.fetch_all(sql)
        data = []
        total_version = []
        reserved_project = ["普晟惠-PD-L1及CD8蛋白表达检测", "普晟惠-MSI微卫星不稳定性检测","普晟畅-结直肠癌靶向用药12基因检测","普晟朗-肺癌靶向用药15基因检测", "普晟和-肿瘤靶向化疗用药83基因检测","普益康-肿瘤个体化诊疗620基因检测"]
        if type == "非ngs":
            filter_type = "noNgs"
        else:
            filter_type = type
        if project_tmp:
            for d in project_tmp:
                if d[1] in reserved_project and filter_type == d[5]:
                    if not version:
                        tmp = {"title": d[1], "text": "", "url": "",'type':d[5], "version":d[3]}
                        data.append(tmp)
                    elif version == d[3]:
                        tmp = {"title": d[1], "text": "", "url": "", 'type': d[5], "version": d[3]}
                        data.append(tmp)
                    if d[3] not in total_version:
                        total_version.append(d[3])
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        return_info = {"success":True, "info": data, "version":total_version}
        resp.body = json.dumps(return_info, ensure_ascii=False)


class GetSampleInfo(object):
    """根据样本的编写获得样本的其他信息"""
    def on_post(self, req, resp):
        data = req.params
        samplename = data["sampleName"]
        db = Mysql(table_name="GLORIA_MYSQL")
        sql = """select * from sample_mx where S_MCODE = \"%s\";""" % samplename
        data = db.fetch_all(sql)
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        if data:
            result = '{"success":true, "info": "样本存在!" }'
        else:
            result = '{"success":false, "info": "样本不存在!" }'
        resp.body = result


class reportInfo(object):
    def on_post(self, req, resp):
        params = req.params
        sample_code = params["sampleCode"]   # 样本编号
        project = params["title"]  # 报告模板
        report_name = params["reportName"]   # 文件名
        user_name = params["userName"]
        workflow_name = params["workflowName"]  # 报告结果名称
        tumorInfiltrating = params["tumorInfiltrating"]
        pd28Tumor = params["pd28Tumor"]
        pd28Lymph = params["pd28Lymph"]
        pd142Tumor = params["pd142Tumor"]
        pd142Lymph = params["pd142Lymph"]
        cd8 = params["cd8"]
        tumorPercent = params["tumorPercent"]
        tumorLevel = params["tumorLevel"]
        msiStable = params["msiStable"]
        # uid= params['uid']
        filenameLeft = params['filenameLeft']
        filenameRight = params['filenameRight']
        uidleft = params['uidLeft']
        uidright = params['uidRight']
        db = Mysql(table_name="REPORT_MYSQL")
        if not uidleft:
            uidLeft = None
        else:
            sql = """select * from png where ID=\"%s\";""" % uidleft
            data = db.fetch_one(sql)
            uidLeft = os.path.join(data[-1], data[-2])
        if not uidright:
            uidRight = None
        else:
            sql = """select * from png where ID=\"%s\";""" % uidright
            data = db.fetch_one(sql)
            uidRight = os.path.join(data[-1], data[-2])
        msi_list = {"NR21":None,"NR24":None,"NR27":None,
                    "BAT25":None,"BAT26":None,"MONO27":None,"PentaC":None,"PentaD":None}
        if msiStable:
            for item in msi_list.keys():
                if item in msiStable.split("_"):
                    msi_list[item] = "稳定"
                else:
                    msi_list[item] = "不稳定"
            print('msi_list', msi_list)
        reportStable = params["reportStable"]

        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"

        mysql = Mysql(table_name="MYSQL")
        sql = "select * from login where userName=\"%s\";" % user_name
        data1 = mysql.fetch_one(sql)
        print('data1', data1)
        if data1:
            usernameid = data1[-1]
            print(usernameid)
            out_path = "/home/khl/web/dist/downloads/tmp"
            if "组织版" in project:
                project_name = project.split('-组织版')[0]
                type = "组织版"
            elif "血液版" in project:
                project_name = project.split('-血液版')[0]
                type = "血液版"
            else:
                project_name = project
                type = None
            resp.status = falcon.HTTP_200
            resp._headers["Access-Control-Allow-Origin"] = "*"

            if type:
                sql = "select * from projectName where projectName=\"%s\" and type = \"%s\";" % (project_name, type)
            else:
                sql = "select * from projectName where projectName=\"%s\";" % project_name
            data = db.fetch_one(sql)
            projectname_id = data[0]
            analysis_name = data[1]
            xml_path = data[8]
            if data:
                print('xml_path', xml_path)

                if project_name == "普晟惠-PD-L1及CD8蛋白表达检测":
                    tumor_infiltrating = get_tumor_infiltrating(tumorInfiltrating, pd28Tumor, pd28Lymph,
                                            pd142Tumor,pd142Lymph, cd8,tumorPercent, tumorLevel)
                    msi_info = None
                if project_name == "普晟惠-MSI微卫星不稳定性检测":
                    msi_info = get_msi_info(NR21=msi_list['NR21'], NR24=msi_list['NR24'],
                                            NR27=msi_list['NR27'], BAT25=msi_list['BAT25'], BAT26=msi_list['BAT26'],
                                            MONO27=msi_list['MONO27'], PentaC=msi_list['PentaC'],
                                            PentaD=msi_list['PentaD'],reportStable=reportStable)
                    print('msi_info', msi_info)
                    tumor_infiltrating = None

                pdfpath, pdfurl = templateword(xml_path, out_path, report_name, sample_code, project_name, usernameid, type, tumor_infiltrating=tumor_infiltrating, msi_info=msi_info, image_left=uidLeft,image_right=uidRight)
                ###ID,templateID,analysisName,sampleName,workflowName,DATETIME,STATUS,path,reportName,pdfPath,pdfUrl,userNameId)
                #sql_tmp = "insert into report(ID,templateID,analysisName,sampleName,workflowName,DATETIME,STATUS,path,reportName,pdfPath,pdfUrl,userNameId) values(UUID(), \"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\")"
                sql_tmp = "insert into report values(UUID(), \"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\",\"%s\")"
                sql = sql_tmp % (
                projectname_id, analysis_name, sample_code, workflow_name, datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "完成", "/home/khl/web/dist/", report_name, pdfpath, pdfurl, usernameid)

                print(sql)
                db.execute(sql)
                db.commit()
                return_info = '{"success":true, "info": "Report generation success!" }'
                resp.body = return_info
            else:
                return_info = '{"success":false, "info": "Report generation failure!" }'
                resp.body = return_info
        else:
            resp.body = '{"success":false, "info": "Report generation failure!" }'

class workflow(object):
    def on_post(self, req, resp):
        params = req.params
        user_name = params["userName"]
        size = int(params["size"])
        page = int(params['page'])
        queryversion = params["version"]
        filter = params["filter"]
        i = 0
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"

        outversion = []  # 列表，嵌套字典
        totalversion = []  # 列表 存储unique version信息
        mysql = Mysql(table_name="MYSQL")
        sql = "select * from login where userName=\"%s\";" % user_name
        data = mysql.fetch_one(sql)
        result = {}
        if data:
            usernameid = data[-1]
            db = Mysql(table_name="REPORT_MYSQL")
            """{title: '分析名称',key: 'analysisName'},
                {title: '样本名称',key: 'sampleName'},
                {title: '流程名称',key: 'flowName'},
                {title: '版本',key: 'version'},
                {title: '日期',key: 'date'},
                {title: '状态',key: 'status'},"""
            sql = "select * from report where userNameId=\"%s\" order by datetime desc;" % usernameid
            info = []
            data = db.fetch_all(sql)
            sql = "select * from projectName;"
            projectName_id_tmp = db.fetch_all(sql)
            projectName_id_list = {}
            for p in projectName_id_tmp:
                projectName_id_list[p[0]] = p
            if data:
                for tmp in data:
                    analysisName=tmp[2]
                    sampleName=tmp[3]
                    flowName=tmp[4]
                    projectid = tmp[1]
                    if projectid in projectName_id_list.keys():
                        version = projectName_id_list[projectid][3]  #添加version信息
                        author = projectName_id_list[projectid][6]   #添加author信息
                        date = tmp[5]
                        status = tmp[6]
                        oldpdfpath = tmp[-3]
                        oldpdfurl = tmp[-2]
                        pdfpath = oldpdfpath.replace('/home/khl/web/dist', "http://192.168.1.144:8096")
                        pdfurl = oldpdfurl.replace('/home/khl/web/dist', "http://192.168.1.144:8096")
                        sample_code = tmp[3]  # 样本编号
                        report_name = os.path.basename(oldpdfpath)
                        if version not in totalversion:
                            tmp_data = {"key": version, "value": version}
                            totalversion.append(version)
                            outversion.append(tmp_data)
                            i += 1
                        if filter and queryversion:
                            if ((filter in analysisName  or filter in sampleName) and version == queryversion):
                                aaaa = {"title": analysisName, "sampleName":sampleName ,"author":author,
                                        "reportName":report_name, "sampleCode":sample_code,
                                        "workflowName": flowName, "version": version, "date": date, "status": status,
                                        "pdfpath": pdfpath, "pdfurl": pdfurl, "uniqueId": tmp[0]}
                                info.append(aaaa)
                        elif filter and not queryversion:
                            if filter in analysisName  or filter in sampleName:
                                aaaa = {"title": analysisName, "sampleName": sampleName,"author":author,
                                        "reportName": report_name, "sampleCode": sample_code,
                                        "workflowName": flowName, "version": version, "date": date, "status": status,
                                        "pdfpath": pdfpath, "pdfurl": pdfurl, "uniqueId": tmp[0]}
                                info.append(aaaa)
                        elif not filter and queryversion:
                            if version == queryversion:
                                aaaa = {"title": analysisName, "sampleName": sampleName, "author":author,
                                        "reportName": report_name, "sampleCode": sample_code,
                                        "workflowName": flowName, "version": version, "date": date, "status": status,
                                        "pdfpath": pdfpath, "pdfurl": pdfurl, "uniqueId": tmp[0]}
                                info.append(aaaa)
                        elif not filter and not queryversion:
                            aaaa = {"title": analysisName, "sampleName": sampleName,"author":author,
                                    "reportName": report_name, "sampleCode": sample_code,
                                    "workflowName": flowName, "version": version, "date": date, "status": status,
                                    "pdfpath": pdfpath, "pdfurl": pdfurl, "uniqueId": tmp[0]}
                            info.append(aaaa)
                    else:
                        result["success"] = False
                        result["info"] = "Report failure!"
                        resp.body = json.dumps(result, ensure_ascii=False)
                        # resp.body = '{"success":false, "info": "Report failure!" }'
                range_tmp = range(len(info))
                mod = len(info) % size  # 取余
                reminder = int(len(info) / size)
                if mod != 0:
                    if page == (reminder + 1):
                        range_id = range_tmp[(size * (page - 1)):len(info)]
                    else:
                        range_id = range_tmp[(size * (page - 1)):(size * (page))]
                else:
                    range_id = range_tmp[(size * (page - 1)):(size * page)]
                out_json = []
                for index in range_id:
                    out_json.append(info[index])
                result["success"] = True
                result["length"] = str(len(out_json))
                result["total"] = str(len(info))
                result["totalVersion"] = outversion
                result["info"] = out_json
                resp.body = json.dumps(result, ensure_ascii=False)
            else:
                result["success"] = True
                result["info"] = []
                resp.body = json.dumps(result, ensure_ascii=False)

class deleteWorkflow(object):
    def on_post(self, req, resp):
        params = req.params
        uniqueid = params["id"]
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        db = Mysql(table_name="REPORT_MYSQL")
        sql = "select * from report where ID=\"%s\";" % uniqueid
        data = db.fetch_one(sql)
        if data:
            sql = "delete from report where ID=\"%s\"" % uniqueid
            try:
                db.execute(sql)
                db.commit()
                result = {"success":True,"info":"删除成功"}
            except Exception:
                result = {"success":False,"info":"删除失败"}
        else:
            result = {"success": True, "info": "删除成功"}
        resp.body = json.dumps(result, ensure_ascii=False)

class analysis(object):
    def on_post(self, req, resp):
        params = req.params
        result = {}
        user_name = params["userName"]
        size = int(params["size"])
        page = int(params['page'])
        queryversion = params["version"]
        filter = params["filter"]
        ngs_type = params["type"]   # 添加过滤参数  all ngs noNgs
        db = Mysql(table_name="REPORT_MYSQL")
        sql = "select * from projectName;"
        outversion = []  # 列表，嵌套字典
        totalversion = []  # 列表 存储unique version信息
        i = 0
        data=db.fetch_all(sql)
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        if data:
            result["success"] = True
            info = []
            for d in data:
                description = d[1]
                application = d[1]
                version = d[3]
                date = d[2]
                author = d[6]
                tmp_url = d[7]
                if tmp_url:
                    url = tmp_url.replace('/home/khl/web/word', 'http://192.168.1.144:8094')
                else:
                    url = None
                if version not in totalversion:
                    if ngs_type == "ngs" or ngs_type == "noNgs":
                        if d[-2] == ngs_type:
                            tmp = {"key":version,"value":version}
                            totalversion.append(version)
                            outversion.append(tmp)
                            i += 1
                        else:
                            continue
                    else:
                        tmp = {"key": version, "value": version}
                        totalversion.append(version)
                        outversion.append(tmp)
                        i += 1
                if filter and queryversion:
                    if filter in application and version == queryversion:
                        if ngs_type == "ngs" or ngs_type == "noNgs":
                            if ngs_type == d[-2]:
                                info_tmp = {"description":description,"application":application,
                                    "version":version,"date":date, 'author': author, 'url':url}
                                info.append(info_tmp)
                            else:
                                continue
                        else:
                            info_tmp = {"description": description, "application": application,
                                        "version": version, "date": date, 'author':author, 'url':url}
                            info.append(info_tmp)
                elif filter and not queryversion:
                    if filter in application:
                        if ngs_type == "ngs" or ngs_type == "noNgs":
                            if ngs_type == d[-2]:
                                info_tmp = {"description": description, "application": application,
                                            "version": version, "date": date, 'author':author, 'url':url}
                                info.append(info_tmp)
                            else:
                                continue
                        else:
                            info_tmp = {"description": description, "application": application,
                                        "version": version, "date": date, 'author':author, 'url':url}
                            info.append(info_tmp)
                elif not filter and queryversion:
                    if version == queryversion:
                        if ngs_type == "ngs" or ngs_type == "noNgs":
                            if ngs_type == d[-2]:
                                info_tmp = {"description": description, "application": application,
                                            "version": version, "date": date, 'author':author, 'url':url}
                                info.append(info_tmp)
                            else:
                                continue
                        else:
                            info_tmp = {"description": description, "application": application,
                                        "version": version, "date": date, 'author':author, 'url':url}
                            info.append(info_tmp)
                elif not filter and not queryversion:
                    if ngs_type == "ngs" or ngs_type == "noNgs":
                        if ngs_type == d[-2]:
                            info_tmp = {"description": description, "application": application,
                                        "version": version, "date": date, 'author':author, 'url':url}
                            info.append(info_tmp)
                        else:
                            continue
                    else:
                        info_tmp = {"description": description, "application": application,
                                    "version": version, "date": date, 'author':author, 'url':url}
                        info.append(info_tmp)
            print('info',info)
            range_tmp = range(len(info))
            mod = len(info) % size  # 取余  8
            reminder = int(len(info) / size)   # 3
            print(reminder)
            print('page', page, 'reminder', reminder)
            if mod != 0:
                if page == (reminder+1):
                    range_id = range_tmp[(size * (page - 1)):len(info)]
                else:
                    range_id = range_tmp[(size * (page - 1)):(size * (page))]
            else:
                range_id = range_tmp[(size * (page - 1)):(size * page)]
            out_json = []
            for index in range_id:
                out_json.append(info[index])
            print('out_json', len(out_json))
            print('range_id',size*page, size*(page+1),range_id)
            result["info"] = out_json
            result["length"] = str(len(out_json))
            result["total"] = str(len(info))
            result["totalVersion"] = outversion
            resp.body = json.dumps(result, ensure_ascii=False)
        else:
            result["success"] = False
            result["info"] = "Report failure!"
            resp.body = json.dumps(result, ensure_ascii=False)

class upload(object):
    def on_post(self, req, resp):
        image = req.get_param('file')
        filename = image.filename
        png_path = '/home/khl/web/upload/%s' % filename
        with open(png_path, 'wb') as temp_file:
            temp_file.write(image.file.read())
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        if os.path.exists(png_path):
            url_path = png_path.replace('/home/khl/web/upload', 'http://192.168.1.144:8095')
            url_info = {"url":url_path}
            return_info = {"success":True, "info": url_info }
        else:
            return_info = {"success": False, "info": "FALSE!"}
        resp.body = json.dumps(return_info, ensure_ascii=False)

class analysisquery(object):
    def on_post(self, req, resp):
        params = req.params
        result = {}
        user_name = params["userName"]
        size = int(params["size"])
        page = int(params['page'])
        queryversion = params["version"]
        filter = params["filter"]
        db = Mysql(table_name="REPORT_MYSQL")
        sql = "select * from projectName;"
        print(sql)
        outversion = []  # 列表，嵌套字典
        totalversion = []  # 列表 存储unique version信息
        i = 0
        data=db.fetch_all(sql)
        resp.status = falcon.HTTP_200
        resp._headers["Access-Control-Allow-Origin"] = "*"
        if data:
            result["success"] = True
            info = []
            for d in data:
                description = d[1]
                application = d[1]
                version = d[3]
                date = d[2]
                if version not in totalversion:
                    tmp = {}
                    tmp[str(i)] = version
                    outversion.append(tmp)
                    i += 1
                if filter and queryversion:
                    if filter in application and version == queryversion:
                        info_tmp = {"description":description,"application":application,
                            "version":version,"date":date}
                        info.append(info_tmp)
                elif filter and not queryversion:
                    if filter in application:
                        info_tmp = {"description": description, "application": application,
                                    "version": version, "date": date}
                        info.append(info_tmp)
                elif not filter and queryversion:
                    if version == queryversion:
                        info_tmp = {"description": description, "application": application,
                                    "version": version, "date": date}
                        info.append(info_tmp)
                elif not filter and not queryversion:
                    info_tmp = {"description": description, "application": application,
                                "version": version, "date": date}
                    info.append(info_tmp)

            range_tmp = range(len(info))
            mod = len(info) % size  # 取余
            reminder = int(len(info) / size)
            if mod != 0:
                if page == 1:
                    range_id = range_tmp[0:size]
                elif page == (reminder - 1):
                    range_id = range_tmp[(size * (page - 1)):len(info)]
                else:
                    range_id = range_tmp[(size*page):mod]
                    #range_id = range_tmp[(size * (page - 1)):(size * page)]
            else:
                range_id = range_tmp[(size * (page - 1)):(size * page)]
            out_json = []
            print(info)
            for index in range_id:
                out_json.append(info[index])
            result["info"] = out_json
            result["length"] = str(len(out_json))
            result["total"] = str(len(info))
            result["totalVersion"] = outversion
            resp.body = json.dumps(result, ensure_ascii=False)
        else:
            result["success"] = False
            result["info"] = "Report failure!"
            resp.body = json.dumps(result, ensure_ascii=False)


app.add_route("/login", Login())
app.add_route("/register", Register())
app.add_route("/sample", query_sample())
app.add_route("/download", Download())
app.add_route("/project", project_name())
app.add_route("/sampleinfo", GetSampleInfo())
app.add_route("/reportinfo", reportInfo())
app.add_route("/upload", upload())
app.add_route("/workflow", workflow())
app.add_route("/delete", deleteWorkflow())
app.add_route("/analysis", analysis())
app.add_route("/analysisquery", analysisquery())


# app.add_route("/samplefilter", query_sample_filter())

# httpd = simple_server.make_server("192.168.1.144", 8099, app)
# httpd.serve_forever()
